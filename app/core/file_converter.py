"""File format converters — PDF, DOCX, XLSX, PNG to/from plain text/data.

All functions return {"ok": bool, "text": str, "error": str}.
Missing optional dependencies produce ok=False with a clear install hint.
"""
from __future__ import annotations


def read_pdf(path: str) -> dict:
    """Extract text from a PDF file using pypdf.

    Returns {"ok", "text", "pages", "error"}.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        return {"ok": False, "text": "", "pages": 0,
                "error": "pypdf not installed — run: pip install pypdf"}
    try:
        reader = PdfReader(path)
        pages = len(reader.pages)
        parts: list[str] = []
        for i, page in enumerate(reader.pages):
            extracted = page.extract_text() or ""
            if extracted.strip():
                parts.append(f"[Page {i + 1}]\n{extracted}")
        text = "\n\n".join(parts)
        return {"ok": True, "text": text, "pages": pages, "error": ""}
    except Exception as exc:
        return {"ok": False, "text": "", "pages": 0, "error": str(exc)}


def write_pdf(path: str, text: str) -> dict:
    """Write plain text to a PDF file using reportlab.

    Returns {"ok", "error"}.
    """
    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.pagesizes import LETTER
    except ImportError:
        return {"ok": False,
                "error": "reportlab not installed — run: pip install reportlab"}
    try:
        doc = SimpleDocTemplate(path, pagesize=LETTER)
        styles = getSampleStyleSheet()
        story = []
        for para in text.split("\n\n"):
            para = para.strip()
            if para:
                story.append(Paragraph(para.replace("\n", "<br/>"),
                                       styles["Normal"]))
        doc.build(story)
        return {"ok": True, "error": ""}
    except Exception as exc:
        return {"ok": False, "error": str(exc)}


def read_docx(path: str) -> dict:
    """Extract text from a .docx file using python-docx.

    Returns {"ok", "text", "error"}.
    """
    try:
        from docx import Document
    except ImportError:
        return {"ok": False, "text": "",
                "error": "python-docx not installed — run: pip install python-docx"}
    try:
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs]
        text = "\n".join(paragraphs)
        return {"ok": True, "text": text, "error": ""}
    except Exception as exc:
        return {"ok": False, "text": "", "error": str(exc)}


def write_docx(path: str, text: str) -> dict:
    """Write plain text paragraphs to a .docx file.

    Returns {"ok", "error"}.
    """
    try:
        from docx import Document
    except ImportError:
        return {"ok": False,
                "error": "python-docx not installed — run: pip install python-docx"}
    try:
        doc = Document()
        for para in text.split("\n"):
            doc.add_paragraph(para)
        doc.save(path)
        return {"ok": True, "error": ""}
    except Exception as exc:
        return {"ok": False, "error": str(exc)}


def read_xlsx(path: str) -> dict:
    """Read an .xlsx spreadsheet.

    Returns {"ok", "sheets": {sheet_name: [[row]]}, "text", "error"}.
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        return {"ok": False, "sheets": {}, "text": "",
                "error": "openpyxl not installed — run: pip install openpyxl"}
    try:
        wb = load_workbook(path, read_only=True, data_only=True)
        sheets: dict[str, list[list]] = {}
        text_parts: list[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows: list[list] = []
            for row in ws.iter_rows(values_only=True):
                rows.append([cell for cell in row])
            sheets[sheet_name] = rows
            # Build a plain-text representation
            sheet_lines = [f"[Sheet: {sheet_name}]"]
            for row in rows:
                sheet_lines.append("\t".join("" if v is None else str(v) for v in row))
            text_parts.append("\n".join(sheet_lines))
        wb.close()
        return {"ok": True, "sheets": sheets, "text": "\n\n".join(text_parts), "error": ""}
    except Exception as exc:
        return {"ok": False, "sheets": {}, "text": "", "error": str(exc)}


def write_xlsx(path: str, data: list[list]) -> dict:
    """Write a 2D list to an .xlsx spreadsheet (single sheet named 'Sheet1').

    Returns {"ok", "error"}.
    """
    try:
        from openpyxl import Workbook
    except ImportError:
        return {"ok": False,
                "error": "openpyxl not installed — run: pip install openpyxl"}
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        for row in data:
            ws.append(row)
        wb.save(path)
        return {"ok": True, "error": ""}
    except Exception as exc:
        return {"ok": False, "error": str(exc)}


def read_png(path: str) -> dict:
    """Read a PNG file.

    Returns {"ok", "width", "height", "mode", "text", "error"}.
    'text' is a human-readable description such as 'PNG image 800x600 RGB'.
    """
    try:
        from PIL import Image
    except ImportError:
        return {"ok": False, "width": 0, "height": 0, "mode": "",
                "text": "", "error": "Pillow not installed — run: pip install Pillow"}
    try:
        with Image.open(path) as img:
            width, height = img.size
            mode = img.mode
        description = f"PNG image {width}x{height} {mode}"
        return {"ok": True, "width": width, "height": height,
                "mode": mode, "text": description, "error": ""}
    except Exception as exc:
        return {"ok": False, "width": 0, "height": 0, "mode": "",
                "text": "", "error": str(exc)}


def write_png(path: str, width: int = 800, height: int = 600,
              color: tuple = (255, 255, 255)) -> dict:
    """Create a solid-color PNG (useful for testing). Requires Pillow.

    Returns {"ok", "error"}.
    """
    try:
        from PIL import Image
    except ImportError:
        return {"ok": False,
                "error": "Pillow not installed — run: pip install Pillow"}
    try:
        img = Image.new("RGB", (width, height), color)
        img.save(path, format="PNG")
        return {"ok": True, "error": ""}
    except Exception as exc:
        return {"ok": False, "error": str(exc)}
